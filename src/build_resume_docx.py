#!/usr/bin/env python3
"""Build Yuanjun (Johnny) Dai's resume as a Word .docx file using python-docx."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(0.50)
section.right_margin  = Inches(0.50)
section.top_margin    = Inches(0.45)
section.bottom_margin = Inches(0.40)

# ── Colours ───────────────────────────────────────────────────────────────────
C_BLACK  = RGBColor(0x0D, 0x0D, 0x0D)
C_DARK   = RGBColor(0x1A, 0x1A, 0x1A)
C_GRAY   = RGBColor(0x55, 0x55, 0x55)
C_BLUE   = RGBColor(0x2A, 0x5A, 0xAA)

FONT_NAME = "Calibri"

# ── Helper: clear default styles ─────────────────────────────────────────────
def _set_para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)

def _run(para, text, size=9, bold=False, italic=False, color=None, underline=False):
    run = para.add_run(text)
    run.font.name      = FONT_NAME
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = color
    return run

def para(align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=11):
    p = doc.add_paragraph()
    p.alignment = align
    _set_para_spacing(p, before=before, after=after, line=line)
    return p

def add_hrule(thickness_pt=0.5):
    """Add a horizontal rule via paragraph border."""
    p = doc.add_paragraph()
    _set_para_spacing(p, before=2, after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(thickness_pt * 8)))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '333333')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def section_header(title):
    """Section title with bottom border rule."""
    p = doc.add_paragraph()
    _set_para_spacing(p, before=6, after=3, line=12)
    r = _run(p, title.upper(), size=9.5, bold=True, color=C_BLACK)
    # Add bottom border to this paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '444444')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def entry_header(title, date_str):
    """Job/project title + right-aligned date on same line using a table."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    # Remove table borders
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement('w:tcBdr')
            for side in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tcBdr.append(el)
            tcPr.append(tcBdr)
    # Remove table spacing
    tblPr = tbl._tbl.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top','left','bottom','right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tblCellMar.append(m)
    tblPr.append(tblCellMar)

    left_cell  = tbl.cell(0, 0)
    right_cell = tbl.cell(0, 1)

    # Title cell
    lp = left_cell.paragraphs[0]
    _set_para_spacing(lp, before=0, after=0, line=11)
    _run(lp, title, size=9.5, bold=True, color=C_BLACK)

    # Date cell (right-aligned)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_spacing(rp, before=0, after=0, line=11)
    _run(rp, date_str, size=8.5, color=C_GRAY)

    # Set column widths
    tbl.columns[0].width = Inches(5.5)
    tbl.columns[1].width = Inches(2.0)
    return tbl

def org_line(text, url=None):
    p = doc.add_paragraph()
    _set_para_spacing(p, before=1, after=2, line=10)
    _run(p, text, size=8.5, color=C_GRAY)
    if url:
        _run(p, "  " + url, size=8, color=C_BLUE)
    return p

def bullet(text_parts, before=0, after=2):
    """
    text_parts: list of (text, bold) tuples OR a plain string.
    """
    p = doc.add_paragraph(style='List Bullet')
    _set_para_spacing(p, before=before, after=after, line=11)
    pf = p.paragraph_format
    pf.left_indent    = Inches(0.18)
    pf.first_line_indent = Inches(-0.13)

    if isinstance(text_parts, str):
        _run(p, text_parts, size=8.7, color=C_DARK)
    else:
        for (text, bold) in text_parts:
            _run(p, text, size=8.7, bold=bold, color=C_DARK)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════════════════════════════════════════
p_name = para(align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=2, line=30)
_run(p_name, "YUANJUN (JOHNNY) DAI", size=22, bold=True, color=C_BLACK)

p_contact = para(before=0, after=4, line=11)
_run(p_contact, "216-269-2394   •   yxd429@case.edu   •   github.com/Johnny-dai-git   •   Cleveland, OH",
     size=8.8, color=C_GRAY)

add_hrule(0.75)

# ═══════════════════════════════════════════════════════════════════════════════
# SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
section_header("Summary")
p_sum = para(before=0, after=4, line=11.5)
_run(p_sum,
     "AI Infrastructure Engineer and PhD Candidate at Case Western Reserve University, specializing in "
     "distributed ML systems, eBPF-based observability, and large-scale LLM deployment. Published ",
     size=8.8, color=C_DARK)
_run(p_sum, "8 research papers (6 first-author)", size=8.8, bold=True, color=C_DARK)
_run(p_sum,
     " at IEEE, ACM, and Springer venues. 5+ years of industry experience at ",
     size=8.8, color=C_DARK)
_run(p_sum, "Cisco", size=8.8, bold=True, color=C_DARK)
_run(p_sum, " and ", size=8.8, color=C_DARK)
_run(p_sum, "Broadcom", size=8.8, bold=True, color=C_DARK)
_run(p_sum, " in high-performance networking and systems optimization.", size=8.8, color=C_DARK)

# ═══════════════════════════════════════════════════════════════════════════════
# EDUCATION
# ═══════════════════════════════════════════════════════════════════════════════
section_header("Education")

entry_header("Ph.D. Candidate in Computer Science (AI Infrastructure)", "08/2019 – Present")
org_line("Case Western Reserve University")

entry_header("M.S. in ECE (Computer Microarchitecture)", "")
org_line("University of Pittsburgh")

# ═══════════════════════════════════════════════════════════════════════════════
# RESEARCH EXPERIENCE
# ═══════════════════════════════════════════════════════════════════════════════
section_header("Research & Systems Engineering Experience")

entry_header("Research Assistant, AI Infrastructure & ML Systems", "")
org_line("Case Western Reserve University")

bullets_research = [
    ([("DYNAMIX – RL-Based Adaptive Batch Scheduling for Distributed ML", True),
      ("  |  Paper / Code", False),
      ("\nDesigned a system-signal-driven RL scheduler (BytePS, PyTorch DDP, TensorFlow) with asynchronous "
       "event-driven control, achieving ", False),
      ("46% end-to-end training speedup", True),
      (" on 32+ nodes across HPC and cloud environments.", False)]),

    ([("PRISM – Priority-Aware Transport for Distributed ML Training", True),
      ("  |  Paper / Code", False),
      ("\nBuilt a selective-reliability transport protocol differentiating critical/non-critical gradient traffic "
       "via importance-based prioritization and dual congestion control; achieved ", False),
      ("8–21% training throughput improvement", True),
      (" under 1% packet loss on 32 GPUs with near-zero convergence loss.", False)]),

    ([("Training-Time Side-Channel Attacks on Distributed ML Systems", True),
      ("  |  Paper / Code", False),
      ("\nInstrumented MXNet BSP training to profile communication and power behavior via Intel RAPL telemetry; achieved ", False),
      (">99% layer-type accuracy and <1.5% tensor-size error", True),
      (" in model reconstruction, exposing actionable MLaaS attack surfaces.", False)]),

    ([("Priority-Based eBPF Network Flow Telemetry for AI / DML Infrastructure", True),
      ("  |  Paper / Code", False),
      ("\nBuilt a priority-aware eBPF telemetry system for real-time in-kernel TCP/UDP monitoring with O(n) "
       "approximate Top-K flow tracking; achieved ", False),
      ("96% Top-K accuracy, 100% priority-packet recall", True),
      (", and <1.1% throughput overhead at 10 Gbps.", False)]),

    ([("High-Resolution eBPF System Resource Profiling for DML Workloads", True),
      ("  |  Paper / Code", False),
      ("\nDeveloped a per-PID eBPF profiler with in-kernel Top-K tracking, delivering ", False),
      ("10–14× finer temporal resolution", True),
      (" than top/htop, ~11 ms latency, and >90% Top-K fidelity on commodity Linux servers.", False)]),

    ([("Scotch – Elastic SDN Overlay for Control-Plane Scalability", True),
      ("  |  Paper / Code", False),
      ("\nDesigned a distributed control-plane overlay using OvS-based traffic offloading and large-flow migration; ", False),
      ("reduced controller congestion by >70%", True),
      (" on a 34-node GENI Clos topology.", False)]),

    ([("Enterprise-Grade LLM Production Deployment Platform on Kubernetes", True),
      ("  |  Code", False),
      ("\nBuilt a K8s microservices platform (Gateway / Router / Worker) with vLLM and TensorRT-LLM, "
       "ArgoCD GitOps, Prometheus/Grafana observability, and CI/CD automation across ", False),
      ("50+ Kubernetes deployments", True),
      (".", False)]),

    ([("Domain-Adapted LLMs for Clinical Risk Adjustment Coding", True),
      ("  |  Paper", False),
      ("\nFine-tuned PubMedBERT and Qwen-7B for HCC coding from EHR notes; achieved ", False),
      ("macro-/micro-F1 of 0.775/0.742 and 0.81/0.802", True),
      (" on MIMIC-IV with imbalance-aware training and per-label calibration.", False)]),

    ([("BiLSTM-Based Side-Channel Modeling for Distributed ML Analysis", True),
      ("  |  Paper", False),
      ("\nDesigned a BiLSTM pipeline aligning CPU/memory energy traces with per-layer timestamps; "
       "reliably distinguishes Conv, FC, and activation layers during distributed ML training using ensemble learning.", False)]),
]

for bp in bullets_research:
    bullet(bp, after=3)

# ═══════════════════════════════════════════════════════════════════════════════
# INDUSTRY EXPERIENCE
# ═══════════════════════════════════════════════════════════════════════════════
section_header("Industry Experience")

entry_header("Software Engineer II", "06/2015 – 08/2016")
org_line("Cisco Systems, Inc.")
bullet([("Developed NPU/ASIC features for a 100 Gbps core router serving hyperscalers (Azure, AWS, Alibaba Cloud) "
         "using Broadcom BCM SDK; contributed to QoS and ECMP modules, tuned forwarding pipelines, ", False),
        ("reducing forwarding latency by ~10%", True),
        (" under heavy load in large-scale data center and HPC environments.", False)], after=4)

entry_header("Applied Research Engineer", "09/2016 – 06/2018")
org_line("Midea Group")
bullet([("Led ML model training and deployment for a smart appliance product; designed data pipeline, applied "
         "transfer learning and ensemble techniques, deployed scalable Docker-based services on Alibaba Cloud. "
         "Product launched at 2018 AWE Show and won the ", False),
        ("Red Dot Design Award", True),
        (".", False)], after=4)

entry_header("Technical Product Manager", "07/2018 – 01/2019")
org_line("Xiaomi")
bullet("Defined system-level architecture for an IoT-enabled smart appliance with PID-based temperature control "
       "and cloud-connected app integration, translating product constraints into deployable system designs.", after=4)

entry_header("Software Engineer Intern", "10/2014 – 05/2015")
org_line("Broadcom (Emulex)")
bullet([("Replaced polling-based data acquisition with DMA + circular-buffer architecture in RTOS pipeline, ", False),
        ("reducing CPU load by 30%", True),
        (" and improving sampling rate and control-loop responsiveness.", False)], after=2)
bullet("Ported a Linux-only management CLI for BladeEngine ASIC-based CNA hardware to Windows Server, "
       "adapting system calls, threading models, and error handling.", after=4)

entry_header("Engineer Intern", "01/2014 – 04/2014")
org_line("Tenova I2S")
bullet([("Re-architected an RTOS data acquisition pipeline using DMA and circular buffers, ", False),
        ("reducing CPU load by 42%", True),
        (" and improving sampling rate and control-loop responsiveness.", False)], after=4)

# ═══════════════════════════════════════════════════════════════════════════════
# ADDITIONAL PROJECTS
# ═══════════════════════════════════════════════════════════════════════════════
section_header("Additional Systems Engineering Experience")

p_proj = para(before=0, after=1, line=11)
_run(p_proj, "16-bit MIPS-Like Pipelined CPU Design and Implementation", size=9.5, bold=True, color=C_BLACK)

p_url = para(before=0, after=2, line=10)
_run(p_url, "github.com/Johnny-dai-git/16-bit-MIPS-Like-Pipelined-CPU-Design-and-Implementation",
     size=8, color=C_BLUE)

bullet([("Designed an end-to-end pipelined CPU from custom ISA through RTL, synthesis, and layout; "
         "5-stage pipeline with hazard handling and branch prediction, achieving ", False),
        ("~3× speedup", True),
        (" over single-cycle design.", False)], after=4)

# ═══════════════════════════════════════════════════════════════════════════════
# PUBLICATIONS
# ═══════════════════════════════════════════════════════════════════════════════
section_header("Publications")
p_pub = para(before=0, after=4, line=11.5)
_run(p_pub, "8 research papers (6 first-author)", size=8.8, bold=True, color=C_DARK)
_run(p_pub,
     " in distributed systems, ML systems, and AI infrastructure, published/accepted at "
     "IEEE IPDPS, IEEE CCNC, IEEE ICNC, ACM TOIT, and Springer SecureComm; 2 manuscripts under review at ACM TOPS.  "
     "Full list: ", size=8.8, color=C_DARK)
_run(p_pub, "scholar.google.com/citations?user=maCNFMAAAAJ", size=8.8, color=C_BLUE)

# ═══════════════════════════════════════════════════════════════════════════════
# SKILLS
# ═══════════════════════════════════════════════════════════════════════════════
section_header("Skills")

skills = [
    ("Languages",       "Python, C/C++, Java, Bash, SQL"),
    ("AI / ML",         "PyTorch (DDP/FSDP), TensorFlow, BytePS, Megatron-LM, NCCL, vLLM, TensorRT-LLM"),
    ("Infrastructure",  "Kubernetes, Docker, ArgoCD, Helm, Kustomize, Slurm, HPC, GitHub Actions CI/CD"),
    ("Observability",   "eBPF, Prometheus, Grafana, DCGM Exporter, Intel RAPL, Kernel Tracing"),
    ("Cloud & Networks","AWS, GCP, SDN (OpenFlow / Ryu), P4, NPU/ASIC (Broadcom BCM SDK)"),
    ("Research",        "Distributed ML Systems, ML Security, Side-Channel Analysis, RL-based Scheduling"),
]

for label, value in skills:
    p_sk = para(before=0, after=2, line=11)
    _run(p_sk, label + ": ", size=8.8, bold=True, color=C_BLACK)
    _run(p_sk, value, size=8.8, color=C_DARK)

# ═══════════════════════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════════════════════
out = "/home/johnny/Desktop/job_hunting/resume/Yuanjun_Dai_Resume.docx"
doc.save(out)
import os
print(f"Saved: {out}  ({os.path.getsize(out)/1024:.1f} KB)")
